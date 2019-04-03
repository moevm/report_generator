#!./venv/bin/python3.6
import sys
import os
import json
import re
import subprocess
import itertools
from docx import Document
import mistune
from pathlib import Path
from PIL import Image
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Inches
from docxtpl import DocxTemplate, RichText

document = None
GIT_REPO = "wiki_dir"
PATH_TO_WIKI = "{}/{}.md"
NAME_REPORT = "report.docx"
LOCAL_REPO = "generated_doc.docx"
TEMPLATE = "template"
SETTINGS_FILE = "settings.json"
COURSE_WORK = "KR"
LAB_WORK = "LR"
PATH_TO_TEMPLATE = "templates/{}.docx"
TYPE_OF_WORK = "type"
NOT_VALID = "not valid file"
MD_EXTENSION = ".md"
DICT_FILENAMES = 'download'
NO_FILE_MESSAGE = "No such file {}"
STYLE = "NORMAL {}"
STANDART_FONT = "Times New Roman"
STANDART_FONT_SIZE = 14
STANDART_PLACE_BEFORE = 0
STANDART_PLACE_AFTER = 0
STANDART_LINE_SPACING = 1.5
STANDART_FONT_QUOTE = "calibri"
ALIGN_CENTRE = "centre"
ALIGN_JUSTIFY = "justify"
ALIGN_LEFT = "left"
UNDER_PICTURE = "Рисунок {}{}"
ATTACHMENT = "Приложение"
PICTURE = "picture"
PAGES = "pages_of_wiki"
STANDART_SIZE_PICTURE = 4
BORDER_OF_PICTURE = 1.5
SPEED_OF_REDUCING_PICTURE = 0.8
MAX_HEAD = 6
DOT = "."
FONT_SIZE_CODE = 10
FONT_CODE = 'Consolas'

M_STUDENT = "Студент"
W_STUDENT = "Студентка"
M_W = "M/W"
MAN = "M"
MAN_OR_WOMAN = "manorgirl"
NUMBER = "number"
CATHEDRA = "cathedra"
DISCIPLINE = "discipline"
THEME = "theme"
GROUP = "group"
NAME_OF_STUDENT = "student"
TEACHER = "teacher"
INIT_DATA = "init_data"
CONTEXT_OF_EXPLANATION = "context_of_explanation"
MIN_PAGES = "min_pages"
DATE_START = "date_start"
DATE_FINISH = "date_finish"
DATE_DEFEND = "date_defend"
ANNOTATION = "annotation"
INTRODUCTION = "introduction"

ERROR_MESSAGE_UNOCONV = "Unoconv error: "
UNOCONC_1ST = "/usr/bin/python3"
UNOCONC_2ND = "/usr/bin/unoconv"
UNOCONC_3RD = "-f"
UNOCONC_4TH = "pdf"

NAME_STYLE = "Mystyle"
SPAN_TEXT = "p.add_run(\"{}\",style=\'{}\')\n"
SPAN_EMPHASIS = "{}.italic = True\n"
SPAN_DOUBLE_EMPHASIS = "{}.bold = True\n"
SPAN_CODE = "p = document.add_paragraph()\np.add_run(\"{}\")\np.style = 'BasicUserQuote'\np.add_run().add_break()\n"
SPAN_LINK = "{} ({})"
SPAN_HRULE = "document.add_page_break()\n"

BLOCK = "block"
BLOCK_MATH = 'block_math'
TOKEN_TYPE = "type"
TOKEN_TEXT = "text"

LIST_ITEM = "p = document.add_paragraph('', style = 'BasicUserList')"
LIST = "{}\np.add_run().add_break()\n"
HEADER = "p = document.add_heading('', {})\n{}"
ADD_PICTURE = "add_picture"
END_STR = ':")\n'
RUN_AND_BREAK = 'p.add_run().add_break()'
ADD_PARAGRAPH = "p = document.add_paragraph()"
TABLE2 = "table = document.add_table(rows={}, cols={}, style = 'BasicUserTable')"
TABLE3 = 'document.add_paragraph().add_run().add_break()\n'
TABLE1 = "table.rows[{}].cells[{}].paragraphs[0]{}\n"
PLUS_STR = "{}{}"

alignment_dict = {'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                  'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
                  'centre': WD_PARAGRAPH_ALIGNMENT.CENTER,
                  'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
                  'left': WD_PARAGRAPH_ALIGNMENT.LEFT}

line_space_dict = {1: WD_LINE_SPACING.SINGLE,
                   2: WD_LINE_SPACING.DOUBLE,
                   1.5: WD_LINE_SPACING.ONE_POINT_FIVE,
                   0: WD_LINE_SPACING.EXACTLY}


class MathBlockGrammar(mistune.BlockGrammar):
    block_math = re.compile(r"^\$\$(.*?)\$\$", re.DOTALL)


class MathBlockLexer(mistune.BlockLexer):
    default_rules = [BLOCK_MATH] + mistune.BlockLexer.default_rules

    def __init__(self, rules=None, **kwargs):
        if rules is None:
            rules = MathBlockGrammar()
        super(MathBlockLexer, self).__init__(rules, **kwargs)

    def parse_block_math(self, m):
        self.tokens.append({TOKEN_TYPE: BLOCK_MATH, TOKEN_TEXT: m.group(1)})


class MarkdownWithMath(mistune.Markdown):
    def __init__(self, renderer, **kwargs):
        kwargs[BLOCK] = MathBlockLexer
        super(MarkdownWithMath, self).__init__(renderer, **kwargs)

    def output_block_math(self):
        return self.renderer.block_math(self.token[TOKEN_TEXT])


class PythonDocxRenderer(mistune.Renderer):
    def __init__(self, **kwds):
        super(PythonDocxRenderer, self).__init__(**kwds)
        self.table_memory = []
        self.img_counter = 0

    def header(self, text, level, raw):
        return HEADER.format(level - 1, text)

    def paragraph(self, text):
        if ADD_PICTURE in text:
            return text
        add_break = '' if text.endswith(END_STR) else RUN_AND_BREAK
        return "{}{}".format('\n'.join((ADD_PARAGRAPH, text, add_break)), '\n')

    def list(self, body, ordered):
        return LIST.format(body)

    def list_item(self, text):
        return '\n'.join((LIST_ITEM, text))

    def table(self, header, body):
        number_cols = header.count('\n') - 2
        number_rows = int(len(self.table_memory) / number_cols)
        cells = [TABLE1.format(i, j, self.table_memory.pop(0)[1:])
                 for i, j in itertools.product(range(number_rows), range(number_cols))]
        tmp = "\n".join([TABLE2.format(number_rows, number_cols)] + cells)
        return PLUS_STR.format(tmp, TABLE3)

    def table_cell(self, content, **flags):
        self.table_memory.append(content)
        return content

    # SPAN LEVEL
    def text(self, text):
        return SPAN_TEXT.format(text, NAME_STYLE)

    def emphasis(self, text):
        return SPAN_EMPHASIS.format(text[:-1])

    def double_emphasis(self, text):
        return SPAN_DOUBLE_EMPHASIS.format(text[:-1])

    def block_code(self, code, language):
        code = code.replace('\n', '\\n')
        return SPAN_CODE.format(code)

    def link(self, link, title, content):
        return SPAN_LINK.format(content, link)

    def hrule(self):
        return SPAN_HRULE


class Dword:

    def __init__(self):
        self.num_of_pictures = 1
        self.number_of_paragraph = 0
        self.name = LOCAL_REPO
        self.download_settings()
        self.choose_path_template()
        #self.make_title()
        #self.doc = Document(self.name)
        self.add_text_from_wiki()
        #self.add_final_part()
        #self.save(self.name)

    def create_styles(self):
        global document
        styles = document.styles
        style = styles.add_style(NAME_STYLE, WD_STYLE_TYPE.CHARACTER)
        style.font.size = Pt(STANDART_FONT_SIZE)
        style.font.name = STANDART_FONT

    def add_text_from_wiki(self):
        global document
        document = Document(os.path.abspath(self.path))
        self.create_styles()
        tmp = []

        for path in self.js_content[PAGES]:
            with open(PATH_TO_WIKI.format(GIT_REPO, path.replace(' ', '-')), 'r', encoding="utf-8") as file:
                tmp.append(file.read())

        renderer = PythonDocxRenderer()

        exec(MarkdownWithMath(renderer=renderer)('\n'.join(tmp)))
        document.save(os.path.abspath(NAME_REPORT))

    def make_title(self):
        doc = DocxTemplate(self.path)
        if self.js_content[M_W] == MAN:
            mw = M_STUDENT
        else:
            mw = W_STUDENT
        content = {
            MAN_OR_WOMAN: RichText(mw),
            NUMBER: RichText(self.js_content[NUMBER]),
            CATHEDRA: RichText(self.js_content[CATHEDRA]),
            DISCIPLINE: RichText(self.js_content[DISCIPLINE]),
            THEME: RichText(self.js_content[THEME]),
            GROUP: RichText(self.js_content[GROUP]),
            NAME_OF_STUDENT: RichText(self.js_content[NAME_OF_STUDENT]),
            TEACHER: RichText(self.js_content[TEACHER]),
            INIT_DATA: RichText(self.js_content[INIT_DATA]),
            CONTEXT_OF_EXPLANATION: RichText(self.js_content[CONTEXT_OF_EXPLANATION]),
            MIN_PAGES: RichText(self.js_content[MIN_PAGES]),
            DATE_START: RichText(self.js_content[DATE_START]),
            DATE_FINISH: RichText(self.js_content[DATE_FINISH]),
            DATE_DEFEND: RichText(self.js_content[DATE_DEFEND]),
            ANNOTATION: RichText(self.js_content[ANNOTATION]),
            INTRODUCTION: RichText(self.js_content[INTRODUCTION])

        }
        doc.render(content)
        doc.save(self.name)

    def add_line(self, line, space_after=STANDART_PLACE_AFTER, set_bold=False, font_name=STANDART_FONT,
                 keep_with_next=False, font_size=STANDART_FONT_SIZE, space_before=STANDART_PLACE_BEFORE,
                 line_spacing=STANDART_LINE_SPACING, align=ALIGN_JUSTIFY, keep_together=True):
        self.number_of_paragraph += 1
        style_name = STYLE.format(self.number_of_paragraph)
        paragraph = document.add_paragraph(line)
        paragraph.style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        font = paragraph.style.font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = bool(set_bold)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = alignment_dict.get(align.lower())
        paragraph_format.space_before = Pt(space_before)
        paragraph_format.space_after = Pt(space_after)
        paragraph_format.keep_with_next = keep_with_next
        paragraph_format.line_spacing_rule = line_space_dict.get(line_spacing)
        paragraph_format.keep_together = keep_together

    def add_final_part(self):
        self.add_page_break()
        self.add_line(ATTACHMENT, set_bold=True, align=ALIGN_CENTRE)
        self.add_code()

    @staticmethod
    def convert_to_pdf(docname):
        try:
            subprocess.check_call(
                [UNOCONC_1ST, UNOCONC_2ND, UNOCONC_3RD, UNOCONC_4TH, docname])
        except subprocess.CalledProcessError as e:
            print(ERROR_MESSAGE_UNOCONV, e)

    def save(self, name=NAME_REPORT):
        document.save(name)

    def add_code(self):
        for filename in self.js_content[DICT_FILENAMES]:
            p = Path(os.getcwd()).rglob(filename)
            for path in p:
                code = NOT_VALID
                with open(path) as file:
                    if file is not None:
                        code = file.read()
                    else:
                        print(NO_FILE_MESSAGE.format(path))

                self.add_line(filename, set_bold=True, align=ALIGN_LEFT)
                self.add_line(code, line_spacing=1, align=ALIGN_LEFT, font_name=FONT_CODE, font_size=FONT_SIZE_CODE)

    def add_page_break(self):
        document.add_page_break()

    def download_settings(self):
        with open(SETTINGS_FILE) as file:
            self.js_content = json.load(file)

    def choose_path_template(self):
        if self.js_content[TYPE_OF_WORK] == COURSE_WORK:
            self.path = PATH_TO_TEMPLATE.format(COURSE_WORK)
        elif self.js_content[TYPE_OF_WORK] == LAB_WORK:
            self.path = PATH_TO_TEMPLATE.format(LAB_WORK)
        else:
            self.path = PATH_TO_TEMPLATE.format(TEMPLATE)
