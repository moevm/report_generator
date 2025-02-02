#!./venv/bin/python3.6
import datetime
import json
import os
import subprocess
from pathlib import Path

import markdown2
import requests
from PIL import Image
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches
from docxtpl import DocxTemplate
from markdown2 import Markdown

from app import ABS_PATH
from github_api import Gengit, LOCAL_REPO as github_local
from markdown2html2word import MyHTMLParser, pre_header, pre_blockquote

OAUTH_PART = "/var/www/report_generator/oauth.txt"
SHELL_COMMAND = "cd {ABS_PATH} && { curl -O {download_url}; cd -; }"

GIT_REPO = ABS_PATH.format("wiki_dir")
PATH_TO_WIKI = "{}/{}.md"
NAME_REPORT = "report.docx"
LOCAL_REPO = ABS_PATH.format("generated_doc.docx")
TEMPLATE = "template"
SETTINGS_FILE = "settings.json"
COURSE_WORK = "KR"
LAB_WORK = "LR"
PATH_TO_TEMPLATE = ABS_PATH.format("word_templates/{}.docx")
TYPE_OF_WORK = "type"
NOT_VALID = "not valid file"
MD_EXTENSION = ".md"
DICT_FILENAMES = 'download'
NO_FILE_MESSAGE = "No such file {}"
STYLE = "NORMAL {}"
STANDART_FONT = "Times New Roman"
STANDART_FONT_SIZE = 14
STANDART_PLACE_BEFORE = 0
STANDART_PLACE_AFTER = 0
STANDART_LINE_SPACING = 1.5
STANDART_FONT_QUOTE = "calibri"
ALIGN_CENTRE = "centre"
ALIGN_JUSTIFY = "justify"
ALIGN_LEFT = "left"
UNDER_PICTURE = "Рисунок {}{}"
ATTACHMENT = "Приложение A"
ARTICLE_SOURCE = "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"
PICTURE = "picture"
PAGES = "pages_of_wiki"
STANDART_SIZE_PICTURE = 4
BORDER_OF_PICTURE = 1.5
SPEED_OF_REDUCING_PICTURE = 0.8
DOT = "."
FONT_SIZE_CODE = 10
FONT_CODE = 'Consolas'

M_STUDENT = "Студент"
W_STUDENT = "Студентка"
M_W = "M/W"
MAN = "M"
MAN_OR_WOMAN = "manorgirl"
NUMBER = "number"
CATHEDRA = "cathedra"
DISCIPLINE = "discipline"
THEME = "theme"
GROUP = "group"
NAME_OF_STUDENT = "student"
TEACHER = "teacher"
INIT_DATA = "init_data"
CONTEXT_OF_EXPLANATION = "context_of_explanation"
MIN_PAGES = "min_pages"
DATE_START = "date_start"
DATE_FINISH = "date_finish"
DATE_DEFEND = "date_defend"
ANNOTATION = "annotation"
EN_ANNOTATION = "en_annotation"
LIST_OF_SOURCE = 'list_of_source'
INTRODUCTION = "introduction"
YEAR = 'year'

ERROR_MESSAGE_CONVERT_TO_PDF = "ERROR PDF "
LIBREOFFICE_CONVERT_DOCX_TO_PDF = "libreoffice --headless --convert-to pdf --outdir {} {}"
#LIBREOFFICE_CONVERT_DOCX_TO_PDF = "libreoffice5.1 --headless --convert-to pdf --outdir {} {}"

STANDART_PT = 6
STANDART_INCHES = 0.5
STANDART_PT_HEADER = 10

BLOCK = "block"
BLOCK_MATH = 'block_math'
TOKEN_TYPE = "type"
TOKEN_TEXT = "text"
EMPTY = " "
DASH = "-"

PLUS_STR = "{}{}"
MAIN_TEXT = "main_text"
CODE_TEXT = "code_text"
FORMAT = "format"
FONT = "font"
SIZE = "size"
TYPE_OF_HEADER = "h{}"
ERROR_STYLE_IN_MD = "В Markdown файле есть стиль, который не поддерживается программой!"
DISTANCE_NUMBER_CODE = " "
NOT_MD_FILES = ['.git', '.', '..']

COMMENTS_PR = "Комментарии из пулл-реквестов"
PR = "pull_request"
OWNER_OF_PR = "owner"
REPO_OF_PR = "repo"
NUMBER_OF_PR = "number_of_pr"

PR_SOURCE_CODE = "Исходный код:"
PR_COMMENTS = "Комментарии:"
PR_DIFFS = "Изменения:"
BAD_URL = "Bad url: {}"

alignment_dict = {'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                  'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
                  'centre': WD_PARAGRAPH_ALIGNMENT.CENTER,
                  'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
                  'left': WD_PARAGRAPH_ALIGNMENT.LEFT
                  }

line_space_dict = {1: WD_LINE_SPACING.SINGLE,
                   2: WD_LINE_SPACING.DOUBLE,
                   1.5: WD_LINE_SPACING.ONE_POINT_FIVE,
                   0: WD_LINE_SPACING.EXACTLY}


class Dword:

    def __init__(self, branch, md=None):
        try:
            self.branch = branch
            self.num_of_pictures = 1
            self.number_of_paragraph = 0
            self.name = LOCAL_REPO
            self.download_settings()
            self.choose_path_template()
            self.make_title()
            self.name_report = '{}_report.docx'.format(self.js_content['group'])

            if self.path:
                self.document = Document(os.path.abspath(self.path))
            else:
                self.document = Document()
            self.update_title_list()
            if not md:
                self.add_text_from_wiki()
            else:
                self.add_text_from_md(md)
            if self.js_content[LIST_OF_SOURCE] != "list_of_source":
                self.add_list_of_source()
            self.add_final_part()
            self.add_comments()

            self.document.save(ABS_PATH.format(self.name_report))
        except ValueError as e:
            raise e
        except Exception as e:
            print(e)
            print('unknown error')  # TODO: need to fix this

        # self.save(self.name)

    def convert_format(self):
        for paragraph in self.document.paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.name = self.js_content[MAIN_TEXT][FONT]
                font.size = Pt(self.js_content[MAIN_TEXT][SIZE])

    def h_w(self, dimension):
        height, width = dimension
        h = w = STANDART_SIZE_PICTURE
        if height > width:
            h *= height / width
            if h / w > BORDER_OF_PICTURE:
                while h / w > BORDER_OF_PICTURE:
                    h *= SPEED_OF_REDUCING_PICTURE
        else:
            w *= width / height
            if w / h > BORDER_OF_PICTURE:
                while w / h > BORDER_OF_PICTURE:
                    w *= SPEED_OF_REDUCING_PICTURE
        return h, w

    def add_picture(self, path):
        try:
            paragraph = self.document.add_paragraph()
            paragraph.paragraph_format.alignment = alignment_dict.get(ALIGN_CENTRE)

            h, w = self.h_w(Image.open(path).size)
            paragraph.add_run().add_picture(path, width=Inches(h), height=Inches(w))
        except Exception as e:
            print(e)
            print("Add pict")

    def add_image_by_url(self, url):
        try:
            req = requests.get(url)
        except requests.exceptions.RequestException:
            print(BAD_URL.format(url))
            return
        filepath = ABS_PATH.format(PICTURE)
        with open(filepath, 'wb') as file:
            file.write(req.content)
        self.add_picture(filepath)

    def add_text_from_wiki(self):
        tmp = []
        try:
            if self.js_content[PAGES]:
                for path in self.js_content[PAGES]:
                    with open(PATH_TO_WIKI.format(GIT_REPO, path.replace(EMPTY, DASH)), 'r', encoding="utf-8") as file:
                        tmp.append(file.read())
            else:
                for filename in os.listdir(GIT_REPO):
                    if filename in NOT_MD_FILES:
                        continue
                    with open(PATH_TO_WIKI.format(GIT_REPO, filename[0:-3]), encoding="utf-8") as file:
                        tmp.append(file.read())
        except FileNotFoundError:
            print('No such md file')
        # "* "
        try:
            pre_header(self.document, self.js_content)
            pre_blockquote(self.document)
            parser_html = MyHTMLParser(self.document, self.js_content)
            markdowner = Markdown(extras=["tables", "cuddled-lists", "smarty-pants", "code-friendly"])
            tmp = tmp[0].split('\n')

            for i in range(len(tmp)):
                curr = tmp[i]
                if len(curr) > 2 and curr[0] == '*' and curr[1] == ' ':
                    curr = curr.replace('* ', '    ', curr.count('* ') - 1)
                tmp[i] = curr

            # print(tmp)
            html = markdowner.convert('\n'.join(tmp))
            # print(html)
            parser_html.feed(html)
        except Exception as e:
            print(e, ERROR_STYLE_IN_MD)
        self.document.save(ABS_PATH.format(self.name_report))

    def make_title(self):
        doc = DocxTemplate(self.path)
        if self.js_content[M_W] == MAN:
            mw = M_STUDENT
        else:
            mw = W_STUDENT
        content = {
            MAN_OR_WOMAN: mw,
            NUMBER: self.js_content[NUMBER],
            CATHEDRA: self.js_content[CATHEDRA],
            DISCIPLINE: self.js_content[DISCIPLINE],
            THEME: self.js_content[THEME],
            GROUP: self.js_content[GROUP],
            NAME_OF_STUDENT: self.js_content[NAME_OF_STUDENT],
            TEACHER: self.js_content[TEACHER],
            INIT_DATA: self.js_content[INIT_DATA],
            CONTEXT_OF_EXPLANATION: self.js_content[CONTEXT_OF_EXPLANATION],

            MIN_PAGES: self.js_content[MIN_PAGES],
            DATE_START: self.js_content[DATE_START],
            DATE_FINISH: self.js_content[DATE_FINISH],
            DATE_DEFEND: self.js_content[DATE_DEFEND],
            ANNOTATION: self.js_content[ANNOTATION],
            EN_ANNOTATION: self.js_content[EN_ANNOTATION],
            INTRODUCTION: self.js_content[INTRODUCTION],
            YEAR: datetime.datetime.now().year
        }
        doc.render(content)
        self.path = self.name
        doc.save(self.name)

    def update_title_list(self):
        for paragraph in self.document.paragraphs:
            font = paragraph.style.font
            runs = paragraph.runs
            for run in runs:
                run.font.name = self.js_content[MAIN_TEXT][FONT]
                run.font.size = Pt(self.js_content[MAIN_TEXT][SIZE])
            font.name = self.js_content[MAIN_TEXT][FONT]
            font.size = Pt(self.js_content[MAIN_TEXT][SIZE])

    def add_line(self, line, space_after=STANDART_PLACE_AFTER, set_bold=False, font_name=STANDART_FONT,
                 keep_with_next=False, font_size=STANDART_FONT_SIZE, space_before=STANDART_PLACE_BEFORE,
                 line_spacing=STANDART_LINE_SPACING, align=ALIGN_JUSTIFY, keep_together=True):
        self.number_of_paragraph += 1
        style_name = STYLE.format(self.number_of_paragraph)
        paragraph = self.document.add_paragraph(line)
        # print('ADD LINE {}'.format(line))
        paragraph.style = self.document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        font = paragraph.style.font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = bool(set_bold)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = alignment_dict.get(align.lower())
        paragraph_format.space_before = Pt(space_before)
        paragraph_format.space_after = Pt(space_after)
        paragraph_format.keep_with_next = keep_with_next
        paragraph_format.line_spacing_rule = line_space_dict.get(line_spacing)
        paragraph_format.keep_together = keep_together

    def add_list_of_source(self):
        #LIST_OF_SOURCE: self.js_content[LIST_OF_SOURCE],
        self.add_page_break()
        self.add_line(ARTICLE_SOURCE, set_bold=True, align=ALIGN_CENTRE)
        list_source = self.js_content[LIST_OF_SOURCE]
        index = 1
        source = '{num}. {src}'
        for elem in list_source.split('\n'):
            elem = source.format(num=index, src=elem)
            self.add_line(elem, line_spacing=1.5, align=ALIGN_LEFT)
            index += 1


    def add_final_part(self):
        print(self.js_content[DICT_FILENAMES])
        if len(self.js_content[DICT_FILENAMES]) > 0:
            print('ADDED')
            self.add_page_break()
            self.add_line(ATTACHMENT, set_bold=True, align=ALIGN_CENTRE)
            self.add_code()

    @staticmethod
    def convert_to_pdf(docname):
        try:
            print(LIBREOFFICE_CONVERT_DOCX_TO_PDF.format(github_local, docname))
            subprocess.call(LIBREOFFICE_CONVERT_DOCX_TO_PDF.format(github_local, docname).split())
        except subprocess.CalledProcessError as e:
            print(ERROR_MESSAGE_CONVERT_TO_PDF, e)

    @staticmethod
    def convert_to_pdf_native(path):
        try:
            print(LIBREOFFICE_CONVERT_DOCX_TO_PDF.format(ABS_PATH[:-3], path))
            subprocess.call(LIBREOFFICE_CONVERT_DOCX_TO_PDF.format(ABS_PATH[:-3], path).split())
        except subprocess.CalledProcessError as e:
            print(ERROR_MESSAGE_CONVERT_TO_PDF, e)

    def save(self, name=NAME_REPORT):
        self.document.save(os.path.abspath(name))

    def number_position(self, _number, code_size):
        max_len = len(str(code_size))
        number = str(_number)
        len_number = len(number)
        return PLUS_STR.format(DISTANCE_NUMBER_CODE * (max_len - len_number), number)

    def add_code(self):
        print(self.js_content[DICT_FILENAMES])

        f = open(OAUTH_PART, 'r')
        oauth = f.read()
        f.close()
        oauth = oauth.strip()
        for filename in self.js_content[DICT_FILENAMES]:
            dir_path = '/'.join(filename.split('/')[:-1])  # Путь до файла
            if dir_path == '':
                dir_path = '/'
            if filename[0] == '.':
                filename = filename[1:]
            auth = 'Authorization'
            token = 'token {token}'.format(token=oauth)
            download_url = "https://raw.githubusercontent.com/{owner}/{repo}/{branch}{path}"
            download_url = download_url.format(owner=self.js_content[PR][OWNER_OF_PR],
                                               repo=self.js_content[PR][REPO_OF_PR], branch=self.branch, path=filename)
            print(download_url)
            response = requests.get(download_url, headers={auth: token})
            if response.status_code == 404:
                raise ValueError('Не удалось получить файл с кодом, проверте название ветки и файла')

            code = response.content.decode('utf-8')
            code.strip()
            self.add_line(filename, set_bold=True, align=ALIGN_LEFT)
            number = 0
            lineArr = code.split('\n')
            for line in lineArr:
                number += 1
                self.add_line(
                    DISTANCE_NUMBER_CODE.join((self.number_position(number, len(lineArr)), line.strip('\n').strip('\r'))),
                    line_spacing=1, align=ALIGN_LEFT, font_name=FONT_CODE, font_size=FONT_SIZE_CODE)


    def add_page_break(self):
        self.document.add_page_break()

    def download_settings(self):
        with open(ABS_PATH.format(SETTINGS_FILE), encoding="utf-8") as file:
            self.js_content = json.load(file)

    def choose_path_template(self):
        if self.js_content[TYPE_OF_WORK] == COURSE_WORK:
            self.path = PATH_TO_TEMPLATE.format(COURSE_WORK)
        elif self.js_content[TYPE_OF_WORK] == LAB_WORK:
            self.path = PATH_TO_TEMPLATE.format(LAB_WORK)
        else:
            self.path = None


    def add_comments(self):
        try:
            if not self.js_content[PR][NUMBER_OF_PR]:
                return
            print(self.js_content[PR][NUMBER_OF_PR])
            git = Gengit(branch=self.branch)
            self.add_page_break()
            self.add_line(COMMENTS_PR, align=ALIGN_CENTRE, set_bold=True, keep_with_next=True)
            comments = git.get_comments(self.js_content[PR][OWNER_OF_PR], self.js_content[PR][REPO_OF_PR],
                                        self.js_content[PR][NUMBER_OF_PR])
            print('ELEMENTS')
            for element in comments:
                source_code = element.body_code
                self.add_line('\n{}'.format(PR_SOURCE_CODE), align=ALIGN_LEFT, set_bold=True, line_spacing=1, keep_with_next=True)
                self.add_line(element.filename, align=ALIGN_LEFT, set_bold=True, line_spacing=1, keep_with_next=True)
                self.add_line(source_code, align=ALIGN_LEFT, line_spacing=1, keep_with_next=True,
                              font_name=FONT_CODE, font_size=FONT_SIZE_CODE)
                self.add_line(PR_COMMENTS, align=ALIGN_LEFT, set_bold=True, line_spacing=1, keep_with_next=True)
                for body_element in element.body_comments:
                    self.add_line("{}:{}".format(body_element[0], body_element[1]), line_spacing=1, keep_with_next=True,
                                  align=ALIGN_LEFT)

            self.add_line('\n{}'.format(PR_DIFFS), align=ALIGN_LEFT, set_bold=True, line_spacing=1, keep_with_next=True)

            f = open(OAUTH_PART, 'r')
            oauth = f.read()
            f.close()
            oauth = oauth.strip()
            filenames = []
            for element in comments:
                alreadyAddedDiff = False
                for filename in filenames:
                    if filename == element.filename:
                        alreadyAddedDiff = True
                if alreadyAddedDiff == True:
                    continue
                auth = 'Authorization'
                token = 'token {token}'.format(token=oauth)
                download_url = "https://api.github.com/repos/{owner}/{repo}/commits/{commit}"
                req = download_url.format(owner=self.js_content[PR][OWNER_OF_PR],
                                                   repo=self.js_content[PR][REPO_OF_PR],
                                                   commit=self.branch)
                response = requests.get(req, headers={auth: token})
                response = response.json()
                diffArr = []
                while (len(response["parents"]) != 0) and (response["sha"] != element.commit):
                    for file in response["files"]:
                        if file["filename"] == element.filename:
                            filenames.append(element.filename)
                            diff_string = response["commit"]["committer"]["date"] + '\n'
                            line_arr = file["patch"].split('\n')
                            print(line_arr)
                            prev_skipped = False
                            for line in range(len(line_arr)):
                                if (line_arr[line][0] == '+') or (line != 0 and line_arr[line-1][0] == '+') or (line != len(line_arr)-1 and line_arr[line+1][0] == '+') or (line_arr[line][0] == '-') or (line != 0 and line_arr[line-1][0] == '-') or (line != len(line_arr)-1 and line_arr[line+1][0] == '-'):
                                    prev_skipped = False
                                    diff_string += line_arr[line] + '\n'
                                elif not prev_skipped:
                                    diff_string += ".......\n"
                                    prev_skipped = True
                            diffArr.append(diff_string)
                    req = download_url.format(owner=self.js_content[PR][OWNER_OF_PR],
                                              repo=self.js_content[PR][REPO_OF_PR],
                                              commit=response["parents"][0]["sha"])
                    response = requests.get(req, headers={auth: token})
                    response = response.json()
                for diff in reversed(diffArr):
                    self.add_line(diff, line_spacing=1, align=ALIGN_LEFT, keep_with_next=True,
                                  font_name=FONT_CODE, font_size=FONT_SIZE_CODE)
                # if element.diff:
                #     self.add_line(element.diff, line_spacing=1, align=ALIGN_LEFT, keep_with_next=True)

        except Exception as e:
            print(e)
            print('Failed add comments')

    def add_text_from_md(self, md):
        try:
            pre_blockquote(self.document)
            parser_html = MyHTMLParser(self.document, self.js_content)
            markdowner = Markdown(extras=["tables", "cuddled-lists", "smarty-pants"])
            html = markdowner.convert(md)
            print(html)
            parser_html.feed(html)
        except:
            print(ERROR_STYLE_IN_MD)
        self.document.save(ABS_PATH.format(self.name_report))

# •
