import requests
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.shared import OxmlElement, qn
from docx.opc.constants import RELATIONSHIP_TYPE
from html.parser import HTMLParser
from markdown2 import Markdown
from app import ABS_PATH

EM = "em"
STRONG = "strong"
H1 = "h1"
H2 = "h2"
H3 = "h3"
H4 = "h4"
H5 = "h5"
H6 = "h6"
UL = "ul"
LI = "li"
P = "p"
A = "a"
IMG = "img"
BLOCKQUOTE = "blockquote"
TABLE = "table"
THEAD = "thead"
TR = "tr"
TH = "th"
TD = "td"
CODE = "code"

STANDART_PT_HEADER = 10
STANDART_PT = 6
STANDART_INCHES = 0.25
FORMAT = "format"
FONT = "font"
SIZE = "size"
TYPE_OF_HEADER = "h{}"
MAIN_TEXT = "main_text"
CODE_TEXT = "code_text"

ERROR = "error in HTMLParser. text: {}"
PICTURE_NAME = "picture"
MARKDOWN_EXTRAS = ["tables", "cuddled-lists", "smarty-pants"]


class MyHTMLParser(HTMLParser):

    def error(self, message):
        print(ERROR.format(message))

    def __init__(self, document, json_setting):
        HTMLParser.__init__(self)
        self.settings = json_setting
        self.isBlockQuote = False
        self.document = document
        self.paragraph = self.document.add_paragraph()
        self.run = self.paragraph.add_run()
        self.h = None
        self.bold = False
        self.italic = False
        self.list_level = 0
        self.hyperlink = None
        self.need_dot_li = False
        self.font = json_setting[MAIN_TEXT][FONT]
        self.normal_font = json_setting[MAIN_TEXT][FONT]
        self.code_font = json_setting[CODE_TEXT][FONT]
        self.size = json_setting[MAIN_TEXT][SIZE]
        self.normal_size = json_setting[MAIN_TEXT][SIZE]
        self.code_size = json_setting[CODE_TEXT][SIZE]
        # TABLE SECTION
        self.table = None
        self.table_row = 0
        self.table_col = 0
        self.table_max_col = 0
        self.thead = []
        self.table_mode = False
        self.table_thead_mode = False
        # END TABLE SECTION

    def handle_starttag(self, tag, attrs):
        if tag == EM:
            self.italic = True
        elif tag == P:
            if not self.list_level:
                self.paragraph = self.document.add_paragraph()
            paragraph_format = self.paragraph.paragraph_format
            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph_format.space_after = Pt(STANDART_PT)
            paragraph_format.first_line_indent = Inches(STANDART_INCHES)
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            if self.isBlockQuote:  # it's here because in html <blockquote><p></p></blockquote>
                paragraph_format.first_line_indent = Inches(0)
                paragraph_format.left_indent = Inches(STANDART_INCHES)
                paragraph_format.space_before = Pt(STANDART_PT)
                paragraph_format.space_after = Pt(STANDART_PT)
                paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif tag == STRONG:
            self.bold = True
        elif tag in [H1, H2, H3, H4, H5, H6]:
            self.h = int(tag[1])
        elif tag == UL:
            self.list_level += 1
        elif tag == LI:
            self.paragraph = self.document.add_paragraph()
            self.need_dot_li = True
            self.font = self.normal_font
            self.size = self.normal_size
        elif tag == IMG:
            url = attrs[0][1]
            picture = requests.get(url).content
            with open(ABS_PATH.format(PICTURE_NAME), 'wb') as file:
                file.write(picture)
            try:
                self.document.add_picture(ABS_PATH.format(PICTURE_NAME), width=Inches(4), height=Inches(3))
            except:
                print('ERROR WITH IMAGE {}'.format(url))
        elif tag == A:
            self.hyperlink = attrs[0][1]
        elif tag == BLOCKQUOTE:
            self.isBlockQuote = True
        # TABLE SECTION
        elif tag == TABLE:
            self.table_mode = True
        elif tag == THEAD:
            self.table_thead_mode = True
        elif tag == TH:
            pass
        elif tag == TR:
            if not self.table_thead_mode:
                self.table.add_row()
        elif tag == TD:
            pass
        # END TABLE SECTION
        elif tag == CODE:
            paragraph_format = self.paragraph.paragraph_format
            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            paragraph_format.first_line_indent = Inches(0)
            self.italic = True
            self.size = self.code_size
            self.font = self.code_font

    def handle_endtag(self, tag):
        if tag == EM:
            self.italic = False
        elif tag == STRONG:
            self.bold = False
        elif tag == UL:
            self.list_level -= 1
        elif tag == BLOCKQUOTE:
            self.isBlockQuote = False
        # TABLE SECTION
        elif tag == TABLE:
            self.table_mode = False
        elif tag == THEAD:
            self.table_thead_mode = False
            self.table = self.document.add_table(1, len(self.thead), 'Table Normal') # Table Grid
            self.table_max_col = len(self.thead)
            for i in range(len(self.thead)):
                self.table.rows[0].cells[i].text = self.thead[i]
            self.thead = []
        elif tag == TH:
            if not self.table_thead_mode:
                self.table_col += 1
            if self.table_col == self.table_max_col:
                self.table_col = 0
        elif tag == TR:
            self.table_row += 1
        elif tag == TD:
            self.table_col += 1
            if self.table_col == self.table_max_col:
                self.table_col = 0
        # END TABLE SECTION
        elif tag == CODE:
            self.italic = False
            self.size = self.normal_size
            self.font = self.normal_font

    def handle_data(self, data: str):
        if data.isspace():
            return
        if self.h:
            p = self.document.add_paragraph(data)
            p.style = self.document.styles['h{}'.format(self.h)]
            self.h = None
        elif self.hyperlink:
            add_hyperlink(self.paragraph, self.hyperlink, data, self.font, self.size, '0000FF', True)
            self.hyperlink = None
        elif self.table_thead_mode:
            self.thead.append(data)
        elif self.table_mode:
            print(self.table_col, self.table_row)
            self.table.rows[self.table_row].cells[self.table_col].text = data
        else:
            if self.list_level > 0:
                paragraph_format = self.paragraph.paragraph_format
                paragraph_format.left_indent = Inches(STANDART_INCHES + STANDART_INCHES * (self.list_level - 1)) #+ 0.1 * self.list_level)
                if self.need_dot_li:
                    self.paragraph.add_run('• ')
                    self.need_dot_li = False

            run = self.paragraph.add_run(data.strip('\n'))
            run.bold = self.bold
            run.italic = self.italic
            run.font.name = self.font
            run.font.size = Pt(self.size)


def get_html():
    with open('./test.md') as file:
        md = file.read()
    markdowner = Markdown(extras=["tables", "cuddled-lists", "smarty-pants"])
    html = markdowner.convert(md)
    with open('file.html', 'w') as file:
       file.write(html)
    return html


def save_document(docx):
    docx.save('report.docx')


def pre_header(document, settings):
    for i in range(6):
        custom_header_style = document.styles.add_style('h{}'.format(i + 1), WD_STYLE_TYPE.PARAGRAPH)
        #custom_header_style.base_style = document.styles['Heading {}'.format(i + 1)]
        custom_header_style.font.rtl = True
        custom_header_style.font.name = settings[FORMAT][TYPE_OF_HEADER.format(i + 1)][FONT]
        custom_header_style.font.size = Pt(settings[FORMAT][TYPE_OF_HEADER.format(i + 1)][SIZE])
        custom_header_style.font.bold = True
        custom_header_style.font.italic = False
        custom_header_style.font.underline = False
        custom_header_style.font.color.rgb = RGBColor.from_string('000000')

        paragraph_format = custom_header_style.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph_format.space_after = Pt(STANDART_PT_HEADER)
        paragraph_format.first_line_indent = Inches(STANDART_INCHES)


def add_hyperlink(paragraph, url, text, font="Times New Roman", size=14, color='0000FF', underline=True):
    # https://github.com/python-openxml/python-docx/issues/74#issuecomment-261169410
    # https://github.com/python-openxml/python-docx/issues/383#issue-220027501

    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    rPr.append(c)

    c = OxmlElement('w:sz')
    c.set(qn('w:val'), str(size*2))  # size x2 and convert to str (20 == 10)
    rPr.append(c)

    c = OxmlElement('w:rFonts')
    c.set(qn('w:ascii'), font)
    c.set(qn('w:eastAsia'), font)
    c.set(qn('w:hAnsi'), font)
    c.set(qn('w:cs'), font)
    rPr.append(c)

    if not underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'none')
        rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def pre_blockquote(docx):
    styles = docx.styles
    style = styles.add_style('MyBlockQuote', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_format = style.paragraph_format
    paragraph_format.left_indent = Inches(STANDART_INCHES)
    paragraph_format.space_before = Pt(10)
    paragraph_format.space_after = Pt(10)
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


if __name__ == '__main__':
    docx = Document()
    html = get_html()
    pre_header(docx)
    pre_blockquote(docx)
    parser = MyHTMLParser(docx)
    parser.feed(html)
    save_document(docx)


'''
# Documentation:

* https://github.com/python-openxml/python-docx/issues/352
* https://python-docx.readthedocs.io/en/latest/user/quickstart.html
* https://docs.python.org/3/library/html.parser.html#html.parser.HTMLParser
'''
